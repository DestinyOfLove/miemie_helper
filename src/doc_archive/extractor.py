"""从不同格式的文件中提取文本内容。"""

import hashlib
from pathlib import Path

import cv2
import numpy as np
import pymupdf
from docx import Document
from PIL import Image
from rapidocr_onnxruntime import RapidOCR

_ocr_engine: RapidOCR | None = None


def get_ocr_engine() -> RapidOCR:
    """延迟初始化 OCR 引擎（单例）。"""
    global _ocr_engine
    if _ocr_engine is None:
        _ocr_engine = RapidOCR()
    return _ocr_engine


def preprocess_image(img_array: np.ndarray) -> np.ndarray:
    """图像预处理：灰度化 + 二值化 + 去噪，提升 OCR 识别率。"""
    if len(img_array.shape) == 3:
        gray = cv2.cvtColor(img_array, cv2.COLOR_BGR2GRAY)
    else:
        gray = img_array
    # 自适应二值化
    binary = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
    )
    # 轻度去噪
    denoised = cv2.fastNlMeansDenoising(binary, h=10)
    return denoised


def ocr_image(img_array: np.ndarray) -> str:
    """对图片数组执行 OCR 识别。"""
    engine = get_ocr_engine()
    result, _ = engine(img_array)
    if not result:
        return ""
    # result 是 list of [box, text, confidence]
    lines = [item[1] for item in result]
    return "\n".join(lines)


def extract_from_image(file_path: Path) -> str:
    """从图片文件提取文本。"""
    img = Image.open(file_path)
    img_array = np.array(img)
    preprocessed = preprocess_image(img_array)
    return ocr_image(preprocessed)


def extract_from_pdf(file_path: Path) -> str:
    """从 PDF 文件提取文本。优先提取嵌入文本，否则 OCR。"""
    doc = pymupdf.open(file_path)
    all_text = []

    for page in doc:
        # 先尝试提取嵌入文本
        text = page.get_text().strip()
        if len(text) > 20:
            all_text.append(text)
        else:
            # 文本太少，可能是扫描件，转图片做 OCR
            pix = page.get_pixmap(dpi=300)
            img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                pix.h, pix.w, pix.n
            )
            if pix.n == 4:  # RGBA -> RGB
                img_array = img_array[:, :, :3]
            preprocessed = preprocess_image(img_array)
            ocr_text = ocr_image(preprocessed)
            if ocr_text:
                all_text.append(ocr_text)

    doc.close()
    return "\n".join(all_text)


def extract_from_docx(file_path: Path) -> str:
    """从 .docx 文件提取文本。"""
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    paragraphs.append(cell_text)
    return "\n".join(paragraphs)


# 支持的文件扩展名映射
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
# .doc 需要 LibreOffice 转换，暂不自动处理，记录后提示用户
DOC_EXTENSIONS = {".doc"}

ALL_EXTENSIONS = IMAGE_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS | DOC_EXTENSIONS


def compute_file_hash(file_path: Path) -> str:
    """计算文件 MD5 哈希值。"""
    md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            md5.update(chunk)
    return md5.hexdigest()


def extract_text(file_path: Path) -> tuple[str, str]:
    """根据文件类型提取文本。返回 (text, method)。"""
    suffix = file_path.suffix.lower()

    if suffix in IMAGE_EXTENSIONS:
        return extract_from_image(file_path), "OCR(图片)"
    elif suffix in PDF_EXTENSIONS:
        return extract_from_pdf(file_path), "PDF"
    elif suffix in DOCX_EXTENSIONS:
        return extract_from_docx(file_path), "Word(.docx)"
    elif suffix in DOC_EXTENSIONS:
        return "", "Word(.doc)需转换"
    else:
        return "", "不支持的格式"
