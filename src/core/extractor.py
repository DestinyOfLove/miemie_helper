"""从不同格式的文件中提取文本内容。"""

import hashlib
import logging
import os
import platform
import shutil
import subprocess
import tempfile
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

import cv2
import numpy as np
import pymupdf
from docx import Document
from PIL import Image
from rapidocr_onnxruntime import RapidOCR

from src.config import (
    DOC_EXTENSIONS,
    DOCX_EXTENSIONS,
    IMAGE_EXTENSIONS,
    OCR_DPI,
    OFD_EXTENSIONS,
    PDF_EXTENSIONS,
    WPS_EXTENSIONS,
)

logger = logging.getLogger(__name__)

_ocr_engine: RapidOCR | None = None


def get_ocr_engine() -> RapidOCR:
    """延迟初始化 OCR 引擎（单例）。"""
    global _ocr_engine
    if _ocr_engine is None:
        _ocr_engine = RapidOCR()
    return _ocr_engine


def preprocess_image(img_array: np.ndarray) -> np.ndarray:
    """图像预处理：灰度化 + 二值化 + 去噪，提升 OCR 识别率。"""
    if len(img_array.shape) == 3:
        gray = cv2.cvtColor(img_array, cv2.COLOR_BGR2GRAY)
    else:
        gray = img_array
    # 自适应二值化
    binary = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
    )
    # 轻度去噪
    denoised = cv2.fastNlMeansDenoising(binary, h=10)
    return denoised


def ocr_image(img_array: np.ndarray) -> str:
    """对图片数组执行 OCR 识别。"""
    engine = get_ocr_engine()
    result, _ = engine(img_array)
    if not result:
        return ""
    # result 是 list of [box, text, confidence]
    lines = [item[1] for item in result]
    return "\n".join(lines)


def extract_from_image(file_path: Path) -> str:
    """从图片文件提取文本。"""
    img = Image.open(file_path)
    img_array = np.array(img)
    preprocessed = preprocess_image(img_array)
    return ocr_image(preprocessed)


def extract_from_pdf(file_path: Path) -> str:
    """从 PDF 文件提取文本。优先提取嵌入文本，否则 OCR。"""
    doc = pymupdf.open(file_path)
    all_text = []

    for page in doc:
        # 先尝试提取嵌入文本
        text = page.get_text().strip()
        if len(text) > 20:
            all_text.append(text)
        else:
            # 文本太少，可能是扫描件，转图片做 OCR
            pix = page.get_pixmap(dpi=OCR_DPI)
            img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                pix.h, pix.w, pix.n
            )
            if pix.n == 4:  # RGBA -> RGB
                img_array = img_array[:, :, :3]
            preprocessed = preprocess_image(img_array)
            ocr_text = ocr_image(preprocessed)
            if ocr_text:
                all_text.append(ocr_text)
            # 显式释放大型图像数据，减少内存峰值
            del img_array, preprocessed
            pix = None

    doc.close()
    return "\n".join(all_text)


def extract_from_docx(file_path: Path) -> str:
    """从 .docx 文件提取文本。"""
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            seen: set[int] = set()
            for cell in row.cells:
                cid = id(cell)
                if cid in seen:
                    continue
                seen.add(cid)
                cell_text = cell.text.strip()
                if cell_text:
                    paragraphs.append(cell_text)
    return "\n".join(paragraphs)


# ---------------------------------------------------------------------------
# LibreOffice 转换（.doc / .wps → .docx → python-docx 提取）
# ---------------------------------------------------------------------------

_soffice_path: str | None = None


def find_soffice() -> str | None:
    """查找 LibreOffice soffice 可执行文件路径。找不到返回 None。

    查找优先级：
    1. 环境变量 LIBREOFFICE_PATH（用户显式指定）
    2. 系统 PATH（shutil.which）
    3. 平台默认安装路径
    """
    global _soffice_path
    if _soffice_path is not None:
        return _soffice_path if _soffice_path else None

    candidates: list[str] = []

    # 1. 环境变量优先
    env_path = os.environ.get("LIBREOFFICE_PATH")
    if env_path:
        candidates.append(env_path)

    # 2. 系统 PATH
    path_found = shutil.which("soffice")
    if path_found:
        candidates.append(path_found)

    # 3. 平台默认路径
    system = platform.system()
    if system == "Darwin":
        candidates.append(
            "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        )
    elif system == "Windows":
        for pf_env in ("PROGRAMFILES", "PROGRAMFILES(X86)"):
            pf_dir = os.environ.get(pf_env)
            if pf_dir:
                candidates.append(
                    str(Path(pf_dir) / "LibreOffice" / "program" / "soffice.exe")
                )

    for candidate in candidates:
        if Path(candidate).is_file():
            _soffice_path = candidate
            return _soffice_path

    _soffice_path = ""  # 标记已搜索过，但未找到
    return None


def extract_via_libreoffice(file_path: Path) -> str:
    """通过 LibreOffice 将 .doc/.wps 转为 .docx，再用 python-docx 提取文本。"""
    soffice = find_soffice()
    if not soffice:
        raise FileNotFoundError(
            "未找到 LibreOffice。请安装 LibreOffice: https://www.libreoffice.org/"
        )

    with tempfile.TemporaryDirectory() as tmpdir:
        cmd = [
            soffice,
            "--headless",
            "--norestore",
            "--convert-to", "docx",
            "--outdir", tmpdir,
            str(file_path),
        ]
        result = subprocess.run(
            cmd, capture_output=True, text=True, timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice 转换失败: {result.stderr.strip()}"
            )

        # 查找输出的 .docx 文件
        converted = list(Path(tmpdir).glob("*.docx"))
        if not converted:
            raise FileNotFoundError(
                f"LibreOffice 转换后未找到 .docx 文件 (目录: {tmpdir})"
            )

        return extract_from_docx(converted[0])


# ---------------------------------------------------------------------------
# OFD 提取（ZIP + XML 解析）
# ---------------------------------------------------------------------------

_OFD_NS = "http://www.ofdspec.org/2016"


def extract_from_ofd(file_path: Path) -> str:
    """从 OFD 文件提取文本。OFD 是 ZIP 包，文本在 Content.xml 的 TextCode 节点。"""
    texts: list[str] = []
    with zipfile.ZipFile(file_path, "r") as z:
        # 找到所有 Content.xml（每页一个）
        content_files = sorted(
            n for n in z.namelist() if n.endswith("Content.xml")
        )
        for cf in content_files:
            with z.open(cf) as f:
                try:
                    root = ET.parse(f).getroot()
                except ET.ParseError:
                    logger.warning("OFD 页面 XML 解析失败: %s / %s", file_path, cf)
                    continue
                for tc in root.iter(f"{{{_OFD_NS}}}TextCode"):
                    if tc.text and tc.text.strip():
                        texts.append(tc.text.strip())

    # 如果 XML 解析没有文本（可能是字体混淆），尝试转 PDF 后 OCR
    if not texts:
        logger.info("OFD 未提取到文本，可能存在字体混淆: %s", file_path)

    return "\n".join(texts)


def compute_file_hash(file_path: Path) -> str:
    """计算文件 MD5 哈希值。"""
    md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            md5.update(chunk)
    return md5.hexdigest()


def extract_text(file_path: Path) -> tuple[str, str]:
    """根据文件类型提取文本。返回 (text, method)。"""
    suffix = file_path.suffix.lower()

    if suffix in IMAGE_EXTENSIONS:
        return extract_from_image(file_path), "OCR(图片)"
    elif suffix in PDF_EXTENSIONS:
        return extract_from_pdf(file_path), "PDF"
    elif suffix in DOCX_EXTENSIONS:
        return extract_from_docx(file_path), "Word(.docx)"
    elif suffix in DOC_EXTENSIONS or suffix in WPS_EXTENSIONS:
        fmt = "Word(.doc)" if suffix in DOC_EXTENSIONS else "WPS"
        try:
            text = extract_via_libreoffice(file_path)
            return text, f"{fmt}→LibreOffice"
        except FileNotFoundError as e:
            logger.warning("LibreOffice 不可用: %s", e)
            return "", f"{fmt}需安装LibreOffice"
        except Exception as e:
            logger.error("LibreOffice 转换失败 %s: %s", file_path.name, e)
            return "", f"{fmt}转换失败"
    elif suffix in OFD_EXTENSIONS:
        text = extract_from_ofd(file_path)
        return text, "OFD"
    else:
        return "", "不支持的格式"
